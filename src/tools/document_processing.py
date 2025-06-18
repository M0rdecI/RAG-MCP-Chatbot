from typing import List, Dict, Optional, Any
import os
from pathlib import Path
import logging
from pydantic import BaseModel
import PyPDF2
from docx import Document as DocxDocument
from rich.progress import Progress, SpinnerColumn, TextColumn
from utils.mcp_schema import MCPContext

class Document(BaseModel):
    file_path: str
    content: str
    metadata: Dict[str, Any]  # Change from Dict[str, str] to Dict[str, Any]

class PDFProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        chunks = []
        start = 0
        text = text.replace('\n', ' ').strip()
        
        while start < len(text):
            # Get chunk with overlap
            end = start + self.chunk_size
            chunk = text[start:end]
            
            # Adjust chunk to end at sentence boundary if possible
            if end < len(text):
                last_period = chunk.rfind('.')
                if last_period != -1:
                    chunk = chunk[:last_period + 1]
                    end = start + last_period + 1
            
            chunks.append(chunk.strip())
            start = end - self.chunk_overlap
        
        return chunks
    
    def process(self, file_path: str) -> List[Optional[Document]]:
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                full_text = ""
                
                # Extract text from all pages
                for page in reader.pages:
                    full_text += page.extract_text() + "\n"
                
                if not full_text.strip():
                    logging.warning(f"No text content extracted from {file_path}")
                    return None
                
                # Split into chunks
                chunks = self._chunk_text(full_text)
                documents = []
                
                # Create a document for each chunk
                for i, chunk in enumerate(chunks):
                    metadata = {
                        'source': os.path.basename(file_path),
                        'type': 'pdf',
                        'chunk': i,
                        'total_chunks': len(chunks),
                        'pages': str(len(reader.pages)),
                        'size': os.path.getsize(file_path),
                        'modified': os.path.getmtime(file_path)
                    }
                    
                    documents.append(Document(
                        file_path=file_path,
                        content=chunk,
                        metadata=metadata
                    ))
                
                return documents
                
        except Exception as e:
            logging.error(f"Error processing PDF {file_path}: {e}")
            return None

class DocxProcessor:
    def process(self, file_path: str) -> Optional[Document]:
        try:
            doc = DocxDocument(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return Document(
                file_path=file_path,
                content=content,
                metadata={
                    'source': os.path.basename(file_path),
                    'type': 'docx'
                }
            )
        except Exception as e:
            logging.error(f"Error processing DOCX {file_path}: {e}")
            return None

class DocumentProcessingTool:
    def __init__(self):
        self.processors = {
            '.pdf': PDFProcessor(),
            '.docx': DocxProcessor()
        }
    
    async def process_directory(self, dir_path: str) -> List[Dict]:
        documents = []
        dir_path = Path(dir_path)
        
        if not dir_path.exists():
            logging.error(f"Directory not found: {dir_path}")
            return []
        
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            transient=True
        ) as progress:
            task = progress.add_task("Processing documents...", total=None)
            
            for file_path in dir_path.rglob("*"):
                if file_path.suffix.lower() in self.processors:
                    try:
                        progress.update(task, description=f"Processing {file_path.name}")
                        docs = self.processors[file_path.suffix.lower()].process(str(file_path))
                        if docs:
                            for doc in docs:
                                doc_dict = doc.model_dump()
                                documents.append(doc_dict)
                            logging.info(f"Successfully processed: {file_path.name}")
                        else:
                            logging.warning(f"Failed to process document: {file_path.name}")
                    except Exception as e:
                        logging.error(f"Error processing {file_path.name}: {str(e)}")
                        continue
            
            progress.update(task, completed=True)
        
        if not documents:
            logging.warning("No documents were successfully processed")
        else:
            logging.info(f"Successfully processed {len(documents)} documents")
        
        return documents
    
    def _extract_relevant_content(self, text: str, query: str) -> str:
        """Extract most relevant content using semantic similarity"""
        # Split into semantic chunks
        chunks = self._semantic_chunking(text)
        
        # Score chunks by relevance to query
        chunk_scores = [(chunk, self._calculate_relevance(chunk, query)) 
                       for chunk in chunks]
        
        # Get most relevant chunks
        relevant_chunks = sorted(chunk_scores, 
                               key=lambda x: x[1], 
                               reverse=True)[:2]
        
        return " ".join(chunk for chunk, _ in relevant_chunks)
    
    def to_mcp_context(self, doc: dict) -> MCPContext:
        return MCPContext(
            context_id=doc['metadata'].get('source', '') + f":{doc['metadata'].get('chunk', 0)}",
            content=doc['content'],
            metadata=doc['metadata']
        )